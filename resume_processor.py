import PyPDF2
from docx import Document
from typing import Dict, Optional
import re

class ResumeProcessor:
    """Process and extract information from resume files"""
    
    @staticmethod
    def extract_text_from_pdf(pdf_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text()
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
        return text
    
    @staticmethod
    def extract_text_from_docx(docx_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = Document(docx_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
        return text
    
    @staticmethod
    def extract_text_from_doc(doc_path: str) -> str:
        """Extract text from DOC file (legacy Word format)"""
        # For .doc files, we'll try to use python-docx which may have limited support
        # or we can use a library like python-docx or convert to text
        try:
            # python-docx handles both .doc and .docx, though .doc support is limited
            doc = Document(doc_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error reading DOC: {str(e)}")
    
    @staticmethod
    def extract_resume_data(resume_text: str) -> Dict:
        """Extract key information from resume text"""
        data = {
            'email': ResumeProcessor._extract_email(resume_text),
            'phone': ResumeProcessor._extract_phone(resume_text),
            'skills': ResumeProcessor._extract_skills(resume_text),
            'education': ResumeProcessor._extract_education(resume_text),
            'experience': ResumeProcessor._extract_experience(resume_text),
        }
        return data
    
    @staticmethod
    def _extract_email(text: str) -> Optional[str]:
        """Extract email address from text"""
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        match = re.search(email_pattern, text)
        return match.group(0) if match else None
    
    @staticmethod
    def _extract_phone(text: str) -> Optional[str]:
        """Extract phone number from text"""
        phone_pattern = r'(?:\+1)?[-.\s]?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}'
        match = re.search(phone_pattern, text)
        return match.group(0) if match else None
    
    @staticmethod
    def _extract_skills(text: str) -> list:
        """Extract skills section from resume"""
        skills = []
        # Look for skills section
        skills_patterns = [
            r'(?:skills|competencies|technical skills)[\s\n:]*([^a-z\n]*?)(?:\n\n|education|experience)',
        ]
        
        for pattern in skills_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                skills_text = match.group(1)
                # Split by common delimiters
                skills = [s.strip() for s in re.split(r'[,•\n]', skills_text) if s.strip()]
                break
        
        return skills[:10]  # Return top 10 skills
    
    @staticmethod
    def _extract_education(text: str) -> list:
        """Extract education section from resume"""
        education = []
        # Look for education section
        education_pattern = r'(?:education|academic)[\s\n:]*([^a-z\n]*?)(?:\n\n|experience|skills|$)'
        match = re.search(education_pattern, text, re.IGNORECASE | re.DOTALL)
        
        if match:
            edu_text = match.group(1)
            # Extract degrees
            degree_pattern = r'((?:B\.?A\.?|B\.?S\.?|M\.?A\.?|M\.?S\.?|M\.?B\.?A\.?|Ph\.?D\.?)[^,\n]*)'
            education = re.findall(degree_pattern, edu_text, re.IGNORECASE)
        
        return education[:5]  # Return top 5 educational qualifications
    
    @staticmethod
    def _extract_experience(text: str) -> list:
        """Extract experience/work history from resume"""
        experience = []
        # Look for experience section
        exp_pattern = r'(?:experience|work history|employment)[\s\n:]*([^a-z\n]*?)(?:\n\n|education|skills|$)'
        match = re.search(exp_pattern, text, re.IGNORECASE | re.DOTALL)
        
        if match:
            exp_text = match.group(1)
            # Split job descriptions
            jobs = re.split(r'\n\n+', exp_text)
            experience = [job.strip() for job in jobs if job.strip()]
        
        return experience[:5]  # Return top 5 experiences
    
    @staticmethod
    def get_improvement_suggestions(resume_data: Dict, resume_text: str) -> Dict:
        """Generate improvement suggestions for the resume"""
        suggestions = {
            'formatting': [
                "Use a clean, consistent font (Arial, Calibri, or Times New Roman) in 10-12pt size",
                "Keep margins between 0.5 and 1 inch on all sides",
                "Use bullet points for easy readability and scanning"
            ],
            'content': [],
            'keywords': []
        }
        
        # Check for missing sections
        text_lower = resume_text.lower()
        
        if 'summary' not in text_lower and 'objective' not in text_lower:
            suggestions['content'].append(
                "Add a professional summary or objective statement at the top of your resume for immediate impact."
            )
        
        if not resume_data.get('skills'):
            suggestions['content'].append(
                "Create a dedicated 'Skills' section highlighting your technical and professional abilities."
            )
        else:
            suggestions['content'].append(
                f"Great! You have {len(resume_data.get('skills', []))} skills listed. Consider organizing them by category."
            )
        
        if len(resume_data.get('experience', [])) == 0:
            suggestions['content'].append(
                "Add a 'Work Experience' section with job titles, companies, dates, and key achievements."
            )
        else:
            suggestions['content'].append(
                "Use action verbs like 'developed', 'managed', 'led', 'improved' to describe your accomplishments."
            )
        
        if not resume_data.get('education'):
            suggestions['content'].append(
                "Add an 'Education' section with your degrees, universities, and graduation dates."
            )
        else:
            suggestions['content'].append(
                f"You have {len(resume_data.get('education', []))} educational qualification(s) listed. Well done!"
            )
        
        # Keyword suggestions based on skills
        it_skills = ['Python', 'Java', 'JavaScript', 'SQL', 'AWS', 'Git', 'React', 'Node.js', 
                     'Machine Learning', 'Data Analysis', 'Leadership', 'Communication', 
                     'Project Management', 'Teamwork']
        resume_skills_str = ' '.join(resume_data.get('skills', []))
        
        suggested_keywords = [skill for skill in it_skills if skill.lower() not in resume_skills_str.lower()]
        if suggested_keywords:
            suggestions['keywords'] = [
                f"'{skill}' - Consider adding if you have experience with it (in-demand skill)"
                for skill in suggested_keywords[:4]
            ]
        else:
            suggestions['keywords'] = [
                "Your skills section looks comprehensive! Make sure they match the job description."
            ]
        
        # Additional formatting suggestions
        if len(resume_text) > 2000:
            suggestions['formatting'].append(
                "Consider condensing your resume - aim for 1-2 pages maximum to keep it concise."
            )
        
        if '@' not in resume_text:
            suggestions['formatting'].append(
                "Make sure email address is clearly visible near the top of the page."
            )
        
        suggestions['formatting'].append(
            "Proofread carefully for spelling and grammar errors - these can hurt your chances!"
        )
        
        return suggestions
